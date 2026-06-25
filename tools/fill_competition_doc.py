from __future__ import annotations

import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TEMPLATE = Path(r"C:\Users\27729\Downloads\技术文档.docx")
BACKUP = Path(r"C:\Users\27729\Downloads\技术文档_原始模板_20260622.docx")
OUTPUT = Path(r"C:\Users\27729\Downloads\技术文档_已填写.docx")
DELIVERY = Path(r"D:\Embedded-agent\Competition-documents\技术文档_已填写.docx")
ASSET_DIR = Path(r"D:\Embedded-agent\Competition-documents\assets")

TITLE = "智境方舟\n基于 RA8P1 的可插拔环境感知与 Agent 协同控制终端"
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE = "F4F8FB"
GREEN = "2E7D32"
ORANGE = "C65911"
GRAY = "666666"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline=BLUE, radius=24, width=4):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill="#17324D", spacing=8):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=fnt,
        fill=fill,
        spacing=spacing,
        align="center",
    )


def arrow(draw: ImageDraw.ImageDraw, start, end, color="#4D6A7F", width=7):
    draw.line([start, end], fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 18
    points = [
        (x2, y2),
        (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55),
        (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55),
    ]
    draw.polygon(points, fill=color)


def create_architecture(path: Path):
    image = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(48, True)
    box_font = font(31, True)
    small_font = font(24)
    draw.text((70, 35), "系统总体架构", font=title_font, fill="#17324D")

    boxes = {
        "sensors": (70, 180, 405, 510),
        "ra": (555, 120, 1125, 620),
        "esp": (1275, 180, 1705, 510),
        "cloud": (930, 735, 1705, 925),
    }
    rounded_box(draw, boxes["sensors"], "#E8F5E9", outline="#2E7D32")
    rounded_box(draw, boxes["ra"], "#E3F2FD", outline="#1565C0")
    rounded_box(draw, boxes["esp"], "#FFF3E0", outline="#EF6C00")
    rounded_box(draw, boxes["cloud"], "#F3E5F5", outline="#7B1FA2")

    center_text(draw, (90, 200, 385, 280), "感知与执行", box_font, fill="#1B5E20")
    center_text(
        draw,
        (95, 290, 380, 475),
        "PCA9548A\nCH0 · AHT20\nCH1 · BH1750\nSG90 / 告警",
        small_font,
    )
    center_text(draw, (585, 145, 1095, 235), "瑞萨 RA8P1 本地安全核心", box_font, fill="#0D47A1")
    center_text(
        draw,
        (600, 250, 1080, 565),
        "传感器驱动与诊断\nplatform_ports / device_registry\nLVGL 320×480 触摸 HMI\nrule_program 状态机\n白名单动作与执行反馈\n断网继续运行已部署规则",
        small_font,
    )
    center_text(draw, (1300, 205, 1680, 285), "ESP32-S3 网络协处理", box_font, fill="#9A4D00")
    center_text(draw, (1310, 300, 1670, 465), "UART 115200 · Wi-Fi\nMQTT · NTP\n状态与事件桥接", small_font)
    center_text(draw, (960, 755, 1670, 815), "Hermes + DeepSeek + Web / QQBot", box_font, fill="#5B176E")
    center_text(
        draw,
        (970, 825, 1660, 900),
        "自然语言理解 → 知识约束 → DSL 校验 → MQTT 下发\n实时数据不足时明确报告不可用",
        small_font,
    )

    arrow(draw, (405, 345), (555, 345))
    arrow(draw, (1125, 345), (1275, 345))
    arrow(draw, (1490, 510), (1490, 735))
    arrow(draw, (1280, 780), (1040, 620))
    draw.text((445, 300), "I2C / PWM", font=small_font, fill="#4D6A7F")
    draw.text((1140, 300), "UART", font=small_font, fill="#4D6A7F")
    draw.text((1510, 600), "MQTT", font=small_font, fill="#4D6A7F")
    image.save(path, quality=95)


def create_flow(path: Path):
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(46, True)
    box_font = font(28, True)
    small_font = font(22)
    draw.text((70, 35), "自然语言规则的安全执行流程", font=title_font, fill="#17324D")
    items = [
        ("用户描述场景", "例如：温度升高时\n逐步打开风门"),
        ("Hermes/DeepSeek 规划", "生成候选\nrule_program.v1"),
        ("云端校验", "设备白名单\n角度/时长范围\n知识约束"),
        ("MQTT + ESP32", "签名消息\n与 UART 转发"),
        ("RA8P1 状态机", "ARMED →\nTRIGGERED → DONE"),
        ("证据回流", "deploy_ack\nexecution_state\ntelemetry"),
    ]
    x = 70
    top = 220
    width = 245
    gap = 42
    for index, (name, desc) in enumerate(items):
        box = (x, top, x + width, top + 390)
        fill = ["#E3F2FD", "#F3E5F5", "#FFF8E1", "#FFF3E0", "#E8F5E9", "#ECEFF1"][index]
        outline = ["#1565C0", "#7B1FA2", "#F9A825", "#EF6C00", "#2E7D32", "#546E7A"][index]
        rounded_box(draw, box, fill, outline=outline, radius=22, width=4)
        center_text(draw, (x + 15, top + 35, x + width - 15, top + 135), name, box_font, fill=outline)
        center_text(draw, (x + 20, top + 155, x + width - 20, top + 350), desc, small_font)
        if index < len(items) - 1:
            arrow(draw, (x + width, top + 195), (x + width + gap - 8, top + 195))
        x += width + gap
    image.save(path, quality=95)


def create_hardware_layout(path: Path):
    image = Image.new("RGB", (1700, 980), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(44, True)
    fnt = font(27, True)
    small = font(22)
    draw.text((65, 35), "原型实物组成与接线示意（非实物照片）", font=title_font, fill="#17324D")

    ra = (640, 260, 1080, 700)
    rounded_box(draw, ra, "#E3F2FD", outline="#1565C0")
    center_text(draw, (680, 300, 1040, 420), "CPKHMI-RA8P1\n主控与扩展板", fnt, fill="#0D47A1")
    center_text(draw, (690, 455, 1030, 620), "LCD / 触摸\nI2C-1 / PWM-0\nUART-BRIDGE", small)

    components = [
        ((80, 150, 450, 360), "4.0 英寸触摸屏\nILI9488 + FT6336", "#E8F5E9", "#2E7D32"),
        ((80, 590, 450, 800), "PCA9548A\nAHT20 + BH1750", "#FFF8E1", "#F9A825"),
        ((1250, 150, 1620, 360), "ESP32-S3\nWi-Fi / MQTT", "#FFF3E0", "#EF6C00"),
        ((1250, 590, 1620, 800), "SG90 风门/遮光机构\n外部 5V，共地", "#FCE4EC", "#C2185B"),
    ]
    for box, text, fill, outline in components:
        rounded_box(draw, box, fill, outline=outline)
        center_text(draw, box, text, fnt, fill=outline)
    arrow(draw, (450, 255), (640, 355))
    arrow(draw, (450, 695), (640, 595))
    arrow(draw, (1080, 355), (1250, 255))
    arrow(draw, (1080, 595), (1250, 695))
    draw.text((500, 260), "显示/触摸", font=small, fill="#4D6A7F")
    draw.text((485, 680), "P309/P306", font=small, fill="#4D6A7F")
    draw.text((1125, 260), "UART", font=small, fill="#4D6A7F")
    draw.text((1115, 680), "P105 PWM", font=small, fill="#4D6A7F")
    draw.text(
        (85, 880),
        "说明：当前资料库未包含可用于提交的正面/45°实物照片，正式提交前应在本页后补入真实装配照片。",
        font=small,
        fill="#8A3B12",
    )
    image.save(path, quality=95)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A7B7C5")
        borders.append(element)


def add_cell_text(cell, text: str, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for index, header in enumerate(headers):
        add_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF", size=9)
        set_cell_shading(table.rows[0].cells[index], BLUE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            add_cell_text(cells[index], str(value), size=8.7)
            if row_index % 2 == 1:
                set_cell_shading(cells[index], PALE)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def clear_body(doc: Document):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc: Document, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.74) if align == WD_ALIGN_PARAGRAPH.JUSTIFY else None
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, 10.5, bold)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(9 if level > 1 else 13)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 13 if level == 2 else 11.5, True, BLUE if level <= 2 else "17324D")
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.42)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("• " + item)
        set_run_font(run, 10.2)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_run_font(run, 9, False, GRAY)


def add_picture(doc: Document, path: Path, width_cm=16.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        add_caption(doc, caption)


def page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "system_architecture.png"
    flow = ASSET_DIR / "safe_execution_flow.png"
    hardware = ASSET_DIR / "hardware_layout.png"
    create_architecture(architecture)
    create_flow(flow)
    create_hardware_layout(hardware)

    if not BACKUP.exists():
        shutil.copy2(TEMPLATE, BACKUP)

    doc = Document(str(BACKUP))
    clear_body(doc)
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(105)
    title_run = cover.add_run(TITLE)
    set_run_font(title_run, 24, True, BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(30)
    subtitle_run = subtitle.add_run("全国大学生嵌入式芯片与系统设计竞赛 · 2026\n瑞萨电子赛题技术文档")
    set_run_font(subtitle_run, 15, True, "17324D")
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(150)
    notice_run = notice.add_run("本文不含参赛身份信息")
    set_run_font(notice_run, 9.5, False, GRAY)

    page_break(doc)
    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本作品面向温室育苗柜、实验室样品柜和小型设备间，设计了一套以瑞萨 RA8P1 为本地安全核心的可插拔环境感知与 Agent 协同控制终端。系统通过 PCA9548A 在单一 I2C 标准口下挂载 AHT20 温湿度传感器和 BH1750 光照传感器，RA8P1 完成驱动、诊断、设备注册、数据新鲜度判断、触摸 HMI、规则状态机及 SG90 执行控制；ESP32-S3 仅承担 UART、Wi-Fi、MQTT 和 NTP 网络协处理。屏幕基于 320×480 ILI9488、FT6336 与 LVGL 构建，可按运行时设备注册表动态生成模块入口，避免固定页面与真实接线不一致。云端采用 Hermes/DeepSeek 理解自然语言，将候选控制意图转换为受限 rule_program.v1，经设备白名单、参数范围和知识约束校验后下发，RA8P1 按 ARMED、TRIGGERED、DONE 状态执行，并通过 deploy_ack、execution_state、status 和 telemetry 形成证据闭环。网络断开时，已经部署的规则仍可在 RA8P1 本地运行；传感数据不新鲜时，屏幕、网页和对话端明确报告不可用而不伪造数值。当前工程已完成温度、湿度、光照采集，动态模块显示，SG90 0-180°控制，Wi-Fi 配网，MQTT 状态同步和自然语言规则下发。作品兼顾物联网、嵌入式人工智能与工业 HMI，具备向农业环境控制、机柜运维和实验设备监测迁移的工程基础。",
    )
    add_picture(doc, architecture, 16.2, "图 1  系统总体架构")

    page_break(doc)
    add_heading(doc, "第一部分  作品概述", 1)
    add_heading(doc, "1.1 功能与特性", 2)
    add_body(
        doc,
        "作品以 RA8P1 为主控，提供多传感器接入、动态设备识别、本地触摸显示、规则控制、网络同步和自然语言交互。I2C-1 通过 PCA9548A 同时连接 AHT20 与 BH1750，系统将物理端口、挂载模块和温度/湿度/光照能力分层建模。320×480 触摸屏显示设备、标准口、实时值、Wi-Fi 与执行状态，并能发起 Wi-Fi 扫描和连接。SG90 可在 0-180°范围执行风门或遮光机构动作。云端 Agent 将用户描述转成白名单 DSL，RA8P1 本地状态机负责最终判断和执行，支持 ACK 与执行状态回流。断网时本地规则继续工作，数据失效时各交互端停止给出伪造的“当前值”。",
    )

    add_heading(doc, "1.2 应用领域", 2)
    add_body(
        doc,
        "首选落地场景为温室育苗柜或实验室样品柜：温湿度异常时驱动风门，光照变化时调整遮光机构，并通过本地屏幕和云端页面向值守人员展示状态。相同架构可迁移到小型设备间、弱电机柜、仓储箱体和教学实验平台。项目适合需要“现场可独立运行、联网后可远程协同、传感器可扩展”的场合。通过替换标准口模块和能力映射，可扩展空气质量、水质、振动、门磁、电流等感知；通过更换执行端，可连接继电器、阀门、照明或告警器。作品不依赖云端完成基本控制，因此更适合网络不稳定或要求现场安全兜底的工业与物联网场景。",
    )

    add_heading(doc, "1.3 主要技术特点", 2)
    add_body(
        doc,
        "一是采用“RA8P1 本地安全核心 + ESP32 网络协处理 + 云端 Agent”的分层架构，主要功能代码运行在瑞萨平台。二是建立 platform_ports、device_registry 与 telemetry.samples 统一状态模型，屏幕、ESP32 和云端消费同一份事实。三是在单物理 I2C 口上通过 PCA9548A 挂载多个模块，同时保留真实物理口语义。四是自然语言不能直接操作 GPIO，而要转换为受限 rule_program.v1，并经过设备白名单、角度、时长和冷却参数校验。五是以 deploy_ack、execution_state 和新鲜度时间戳形成可追溯证据链。六是针对嵌入式资源约束采用精简中文字体、轻量竖屏布局和整数安全解析。",
    )

    add_heading(doc, "1.4 主要性能指标", 2)
    add_table(
        doc,
        ["项目", "指标", "依据/说明"],
        [
            ("显示触摸", "320×480；FT6336 原生坐标", "当前配置"),
            ("串口", "115200 bit/s，8N1", "RA8P1-ESP32"),
            ("感知", "温/湿度/光照", "AHT20/BH1750"),
            ("执行", "SG90 0-180°；20 ms", "板端驱动"),
            ("上报", "15 s；事件即时", "ESP32"),
            ("冷启动", "20 s 内数据有效", "验收口径"),
            ("状态", "ACK + ARMED/TRIGGERED/DONE", "实机记录"),
        ],
        [3.2, 5.4, 7.0],
    )

    add_heading(doc, "1.5 主要创新点", 2)
    add_bullets(
        doc,
        [
            "单物理标准口承载多模块，并保持“总线-设备-能力”三层事实一致。",
            "触摸 HMI 从运行时设备注册表动态生成，不写死外设页面。",
            "大模型只生成候选 DSL，RA8P1 白名单状态机掌握最终执行权。",
            "断网可继续运行已部署规则；数据不新鲜时多端一致拒绝伪造。",
            "ACK、执行状态和遥测构成可回放的端到端证据链。",
        ],
    )

    add_heading(doc, "1.6 设计流程", 2)
    add_body(
        doc,
        "项目先打通 RA8P1 显示与触摸，再建立 ESP32 UART/MQTT 桥接；随后接入 AHT20、SG90 和 ACK，扩展为 PCA9548A 下的 AHT20/BH1750 多模块模型；最后将自然语言约束为 rule_program.v1。各阶段均执行“修改-构建-烧录-状态回查-复盘”闭环，避免只凭界面判断成功。",
    )
    add_picture(doc, flow, 16.3, "图 2  自然语言规则的安全执行流程")

    page_break(doc)
    add_heading(doc, "第二部分  系统组成及功能说明", 1)
    add_heading(doc, "2.1 整体介绍", 2)
    add_body(
        doc,
        "系统由感知层、RA8P1 本地核心、网络桥接层和云端协同层组成。感知层通过 PCA9548A 将 AHT20 和 BH1750 接入 I2C-1；RA8P1 将扫描结果、模块身份、能力和样本写入统一状态模型，并驱动 LVGL HMI 与 SG90。ESP32-S3 将 UART 状态转换为 MQTT status、telemetry 和 event，同时接收云端 script。云端 Hermes/DeepSeek 负责自然语言理解，FastAPI 服务负责知识校验、签名、发布和 ACK 等待。执行结果沿原路径返回，网页和对话端只基于最新证据组织回复。",
    )
    add_picture(doc, architecture, 16.3, "图 3  各模块关系")

    add_heading(doc, "2.2 硬件系统介绍", 2)
    add_heading(doc, "2.2.1 硬件整体介绍", 3)
    add_body(
        doc,
        "主控采用 CPKHMI-RA8P1 核心板及扩展板。显示采用 4.0 英寸 ILI9488 IPS 屏，触摸控制器为 FT6336。环境感知采用 AHT20 和 BH1750，并通过 PCA9548A 隔离到不同子通道，避免地址与总线组织冲突。ESP32-S3 通过 SCI0 与 RA8P1 连接。SG90 由 P105 输出控制信号，使用外部 5 V 电源并与系统共地。各传感器使用 3.3 V 供电。",
    )
    add_picture(doc, hardware, 15.8, "图 4  原型实物组成与接线示意")

    add_heading(doc, "2.2.2 机械设计介绍", 3)
    add_body(
        doc,
        "当前版本为功能验证原型，尚未制作定型外壳。建议采用竖屏前面板、主控与 ESP32 后置、传感器通风窗外露、舵机与风门机构侧置的模块化结构。AHT20 应避开主控和稳压器热源，BH1750 应朝向被测光环境并避免屏幕背光直射；SG90 负载机构需要限位和独立供电。正式参赛样机可采用激光切割亚克力或 3D 打印外壳，并在外壳上保留 I2C、PWM、UART 和电源标识。",
    )

    add_heading(doc, "2.2.3 电路各模块介绍", 3)
    add_table(
        doc,
        ["模块", "关键连接", "功能与约束"],
        [
            ("显示屏", "RA8P1 GPIO/SPI → ILI9488", "320×480 显示，LVGL 刷新"),
            ("触摸", "FT6336 → 专用 I2C", "原生 320×480 坐标，不与传感器总线混用"),
            ("PCA9548A", "P309=SDA，P306=SCL", "I2C-1 复用器，地址 0x70"),
            ("AHT20", "CH0，地址 0x38", "温湿度、CRC 和 ACK 诊断"),
            ("BH1750", "CH1，地址 0x23/0x5C", "光照 lux"),
            ("ESP32-S3", "P603→GPIO44，P602←GPIO43", "UART 115200，Wi-Fi/MQTT/NTP"),
            ("SG90", "P105 信号，外部 5 V", "0-180°，必须共地"),
        ],
        [3.1, 6.1, 6.4],
    )

    add_heading(doc, "2.3 软件系统介绍", 2)
    add_heading(doc, "2.3.1 软件整体介绍", 3)
    add_body(
        doc,
        "RA8P1 软件采用 e2 studio/FSP 裸机工程，主循环以 5 ms 节拍驱动端口、传感器、触摸、UI 和控制状态机。底层包括 LCD、FT6336、软件 I2C、AHT20、BH1750、PCA9548A、UART 和 SG90；中间层由 platform_ports 与 device_registry 统一表示端口、模块、能力和样本；应用层由 app_ui 与 esp32_link 完成人机交互和协议处理。ESP32 使用 Arduino 框架，将 UART 行协议映射为 MQTT JSON。云端使用 FastAPI、Hermes/DeepSeek 和持久化状态存储，提供 Web、QQBot、设备状态与部署接口。",
    )

    add_heading(doc, "2.3.2 软件各模块介绍", 3)
    add_table(
        doc,
        ["模块/函数", "关键输入", "关键输出"],
        [
            ("i2c_bus_s1 / aht20 / bh1750", "总线、电源、传感器响应", "诊断、温湿度、lux"),
            ("platform_ports", "扫描结果与采样", "端口、模块、能力、样本"),
            ("device_registry", "platform_ports 状态", "动态设备列表"),
            ("app_ui", "注册表、Wi-Fi、时间、执行状态", "320×480 触摸界面"),
            ("esp32_link", "UART 命令与云端脚本", "ACK、执行状态、状态文本"),
            ("rule_program 状态机", "阈值、动作序列、冷却时间", "ARMED/TRIGGERED/DONE"),
            ("ESP32 bridge", "UART/MQTT/NTP", "status/telemetry/event/script"),
            ("Hermes/DeepSeek", "自然语言与实时上下文", "经约束的查询或控制计划"),
        ],
        [4.3, 5.8, 5.8],
    )
    add_picture(doc, flow, 16.2, "图 5  软件控制流程")

    page_break(doc)
    add_heading(doc, "第三部分  完成情况及性能参数", 1)
    add_heading(doc, "3.1 整体介绍", 2)
    add_body(
        doc,
        "当前已形成可编译、可烧录、可联调的 RA8P1 与 ESP32 双固件工程，以及可运行的云端服务。板端已完成显示、触摸、温湿度、光照、动态模块注册、Wi-Fi 配网、NTP 时间、MQTT 状态与 SG90 控制。项目资料中暂未包含可用于正式提交的正面和 45°实物照片，因此本文采用组成示意图说明；提交前必须补入真实样机照片，不能以生成图替代。",
    )
    add_picture(doc, hardware, 15.8, "图 6  当前原型组成示意；提交前补充真实实物照片")

    add_heading(doc, "3.2 工程成果", 2)
    add_heading(doc, "3.2.1 机械成果", 3)
    add_body(
        doc,
        "已完成主控、屏幕、传感器、ESP32 与舵机的功能原型组织和供电约束设计；最终外壳、风门连杆和传感器通风结构尚需在参赛样机阶段定型。机械设计重点是屏幕可视、传感器不受板卡热源干扰、SG90 独立供电和所有模块可靠固定。",
    )
    add_heading(doc, "3.2.2 电路成果", 3)
    add_body(
        doc,
        "已完成 RA8P1 屏幕与触摸链路、SCI0 与 ESP32-S3 通信、I2C-1/PCA9548A/AHT20/BH1750 多模块链路及 P105/SG90 控制链路。软件可区分 SDA/SCL 卡死、地址 NACK、读取失败、校准失败和 CRC 等故障。当前工程保留 I2C-2 为预留口，不把复用器子通道伪装为额外物理接口。",
    )
    add_heading(doc, "3.2.3 软件成果", 3)
    add_body(
        doc,
        "已实现动态设备列表、设备详情、标准口、Wi-Fi 选择、连续时钟和执行状态页面；ESP32 可上报 ports、samples、status、telemetry 和 event；云端 Web 与 QQBot 可基于新鲜环境数据回答查询，并通过 Hermes/DeepSeek 生成受限规则。历史实机记录证明 rule_program 可完成 TRIGGERED→DONE，传统 threshold_control 路径保持兼容。",
    )

    add_heading(doc, "3.3 特性成果与性能参数", 2)
    add_table(
        doc,
        ["测试项", "当前结果", "结论"],
        [
            ("RA8P1 工程构建", "生成 ELF 与 SREC", "已完成"),
            ("显示/触摸", "320×480 UI 与原生触摸坐标", "已完成"),
            ("多传感器", "PCA9548A + AHT20 + BH1750", "已完成"),
            ("动态模块", "按设备注册表生成入口", "已完成"),
            ("网络桥接", "UART/Wi-Fi/MQTT/NTP", "已完成"),
            ("规则执行", "ACK + ARMED/TRIGGERED/DONE", "已完成"),
            ("断网本地运行", "已部署规则由 RA8P1 状态机执行", "架构支持"),
            ("TinyML 异常检测", "已有采样与训练方案", "待实机部署"),
            ("蜂鸣器", "能力模型已预留", "答辩前实机复测"),
            ("正式统计", "冷启动、丢包率、响应时延多轮数据", "待补测"),
        ],
        [4.3, 7.6, 3.2],
    )

    page_break(doc)
    add_heading(doc, "第四部分  总结", 1)
    add_body(
        doc,
        "本作品以真实硬件闭环为核心，将 RA8P1 的高性能处理、HMI 和丰富外设能力与 ESP32 网络连接、云端 Agent 协同结合起来。项目解决的关键问题不是单一传感器读数，而是如何让物理端口、挂载模块、能力样本、屏幕显示、云端状态和自然语言控制保持一致。通过受限 DSL 与板端状态机，大模型的灵活性被限制在可验证范围内；通过数据新鲜度和证据链，系统避免了离线时仍报告虚假“当前值”的问题。当前原型已具备环境感知、动态显示、远程同步和执行控制基础，可作为农业、机柜和实验设备智能化的通用终端。",
    )
    add_heading(doc, "4.1 可扩展之处", 2)
    add_body(
        doc,
        "后续首先补充外壳、真实风门机构、位置反馈和完整实物照片，并用多轮测试统计冷启动、网络恢复、控制响应和连续运行稳定性。感知侧可扩展 CO₂、VOC、土壤湿度、振动和电流模块；执行侧可扩展继电器、阀门、照明和多路 PWM。算法侧可按现有采样方案训练轻量异常检测模型，利用 RA8P1 Cortex-M85/Helium 在本地完成异常评分，使云端失联时仍能主动告警。协议侧可将 rule_program 扩展为版本化能力图，并增加程序签名、回滚和权限分级。机械侧可设计标准化模块插槽和传感器通风结构，提升部署一致性。",
    )
    add_heading(doc, "4.2 心得体会", 2)
    add_body(
        doc,
        "研发过程中最重要的经验是：嵌入式系统不能只看“界面像不像成功”，必须建立从硬件电气、驱动诊断、板端状态、网络消息到云端页面的连续证据。早期屏幕黑屏并不是 LCD 引脚错误，而是 LVGL 首次刷新调用链超出默认主栈；触摸卡片错位也不是简单扩大点击区就能解决，而是旧坐标偏移与新界面坐标系不一致。AHT20 出现占位符时，问题来自目标端浮点 scanf 行为，最终改为整数安全拆分。多传感器接入后，我们又认识到“一个物理口可挂多个模块”不能在页面上被错误解释成多个物理口，因此建立了端口、模块和能力三层模型。云端 Agent 的调试进一步说明，大模型能够帮助理解自然语言，但不能替代确定性的安全边界；只有经过白名单、参数范围和板端状态机，控制才可复现、可回滚。工程也经历了文档和代码多份镜像漂移的问题，最终确定以真实 RA8P1 工作区为唯一事实源。整个过程让我们从“把模块接起来”转向“让系统在故障、断网和数据缺失时仍然诚实、可解释、可维护”，这也是作品最有价值的工程收获。",
    )

    page_break(doc)
    add_heading(doc, "第五部分  参考文献", 1)
    references = [
        "[1] 全国大学生嵌入式芯片与系统设计竞赛. 2026 芯片应用赛道选题指南：瑞萨电子赛题.",
        "[2] Renesas Electronics. RA8P1 Group Datasheet.",
        "[3] Renesas Electronics. RA Flexible Software Package Documentation.",
        "[4] Renesas Electronics. CPKHMI-RA8P1 Example Projects and Board Documentation.",
        "[5] LVGL LLC. LVGL 9 Documentation.",
        "[6] Ilitek. ILI9488 Preliminary Datasheet.",
        "[7] FocalTech. FT6336U Capacitive Touch Panel Controller Datasheet.",
        "[8] Aosong Electronics. AHT20 Temperature and Humidity Sensor Datasheet.",
        "[9] ROHM. BH1750FVI Ambient Light Sensor Datasheet.",
        "[10] Texas Instruments. TCA9548A Low-Voltage 8-Channel I2C Switch Datasheet.",
        "[11] Espressif Systems. ESP32-S3 Series Datasheet.",
        "[12] OASIS. MQTT Version 3.1.1 Specification.",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.hanging_indent = Cm(0.7)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(reference)
        set_run_font(run, 9.5)

    core_props = doc.core_properties
    core_props.title = TITLE.replace("\n", " - ")
    core_props.subject = "2026 瑞萨电子赛题技术文档"
    core_props.keywords = "RA8P1, 物联网, 嵌入式人工智能, 工业HMI, Agent"
    core_props.comments = "基于真实工程源码与赛题要求生成；实物照片需在提交前补充。"

    doc.save(str(OUTPUT))
    shutil.copy2(OUTPUT, DELIVERY)
    try:
        shutil.copy2(OUTPUT, TEMPLATE)
        print(f"updated_original={TEMPLATE}")
    except PermissionError:
        print(f"original_locked={TEMPLATE}")
    print(f"saved={OUTPUT}")
    print(f"delivery={DELIVERY}")
    print(f"backup={BACKUP}")


if __name__ == "__main__":
    build_document()
